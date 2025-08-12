from typing import List, Dict, Optional
import os, re, hashlib, io
from pdf2image import convert_from_bytes
from PIL import Image
import pytesseract
from concurrent.futures import ThreadPoolExecutor

# Optional imports for non-PDF formats
try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    from bs4 import BeautifulSoup  # beautifulsoup4
except Exception:
    BeautifulSoup = None

try:
    import pandas as pd  # pandas for CSV/XLSX
except Exception:
    pd = None

import sqlite3

# --- Tesseract / Poppler paths (unchanged) -----------------
poppler_path = r"C:\Program Files\poppler-24.08.0\Library\bin"
if os.path.exists(poppler_path):
    os.environ["PATH"] += os.pathsep + poppler_path
tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if os.path.exists(tesseract_path):
    pytesseract.pytesseract.tesseract_cmd = tesseract_path

# --- Helpers ---------------------------------------------
def enhance_image(img: Image.Image) -> Image.Image:
    return img.convert('L').point(lambda x: 0 if x < 140 else 255, '1')


def is_table_block(text_block: str) -> bool:
    lines = text_block.strip().splitlines()
    if len(lines) < 2:
        return False
    avg_pipes = sum(line.count('|') for line in lines) / len(lines)
    return avg_pipes >= 1.5 or any(re.search(r'\s{3,}', ln) for ln in lines)

def ocr_text_to_markdown(text: str) -> str:
    lines = text.strip().splitlines()
    if not lines:
        return text
    header = [c.strip() or "—" for c in re.split(r'\s{2,}|\|', lines[0])]
    if len(header) < 2:
        return text
    rows = [ "| " + " | ".join([c.strip() or "—" for c in re.split(r'\s{2,}|\|', ln)]) + " |"
             for ln in lines[1:] if ln.strip()]
    return "\n".join(["| " + " | ".join(header) + " |",
                      "| " + " | ".join("---" for _ in header) + " |"] + rows)

def ocr_page(idx_img_lang):
    idx, img, lang = idx_img_lang
    try:
        text = pytesseract.image_to_string(enhance_image(img), lang=lang)
        blocks = text.strip().split('\n\n')
        page_text = "\n\n".join(
            ocr_text_to_markdown(b) if is_table_block(b) else b.strip()
            for b in blocks
        )
        page_hash = hashlib.md5(page_text.encode()).hexdigest()
        return {
            "page": idx + 1,
            "text": page_text,
            "page_hash": page_hash
        }
    except Exception as e:
        return {"page": idx + 1, "text": "", "page_hash": ""}

def extract_text_from_pdf(pdf_bytes: bytes, lang: str = 'ben+eng', dpi: int = 300) -> List[Dict]:
    pages = convert_from_bytes(pdf_bytes, dpi=dpi)
    with ThreadPoolExecutor() as ex:
        return list(ex.map(ocr_page, [(i, p, lang) for i, p in enumerate(pages)]))


# ---- Generic loaders for other file types -----------------------------------

def _page_record(text: str, page_number: int = 1) -> Dict:
    text = (text or "").strip()
    return {
        "page": page_number,
        "text": text,
        "page_hash": hashlib.md5(text.encode()).hexdigest() if text else "",
    }


def _extract_text_from_txt(data: bytes, encoding_candidates: Optional[List[str]] = None) -> List[Dict]:
    if encoding_candidates is None:
        encoding_candidates = ["utf-8", "utf-16", "latin-1"]
    text: str = ""
    for enc in encoding_candidates:
        try:
            text = data.decode(enc)
            break
        except Exception:
            continue
    if not text:
        text = data.decode("utf-8", errors="ignore")
    return [_page_record(text, 1)]


def _extract_text_from_docx(data: bytes) -> List[Dict]:
    if docx is None:
        return [_page_record("")]
    try:
        file_like = io.BytesIO(data)
        document = docx.Document(file_like)
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
        # Extract tables as simple markdown
        for tbl in document.tables:
            for row in tbl.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    paragraphs.append("| " + " | ".join(cells) + " |")
        text = "\n".join(paragraphs)
        return [_page_record(text, 1)]
    except Exception:
        return [_page_record("")]


def _extract_text_from_html(data: bytes) -> List[Dict]:
    if BeautifulSoup is None:
        return _extract_text_from_txt(data)
    try:
        html = data.decode("utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        # Convert tables to simple markdown-ish rows
        tables_text: List[str] = []
        for table in soup.find_all("table"):
            for row in table.find_all("tr"):
                cells = [c.get_text(strip=True) for c in row.find_all(["td", "th"])]
                if any(cells):
                    tables_text.append("| " + " | ".join(cells) + " |")
            table.decompose()
        # Remaining text
        for s in soup(["script", "style"]):
            s.decompose()
        text = soup.get_text(" ")
        combined = (text.strip() + ("\n" + "\n".join(tables_text) if tables_text else "")).strip()
        return [_page_record(combined, 1)]
    except Exception:
        return _extract_text_from_txt(data)


def _extract_text_from_csv(data: bytes) -> List[Dict]:
    # If pandas is available, use it to parse and generate structured text via sqlite3 schema
    if pd is not None:
        try:
            text_io = io.StringIO(data.decode("utf-8", errors="ignore"))
            df = pd.read_csv(text_io)
            return _extract_pages_from_dataframe(df, table_name="csv_data")
        except Exception:
            pass
    # Fallback to raw text
    try:
        text = data.decode("utf-8", errors="ignore")
        return [_page_record(text, 1)]
    except Exception:
        return _extract_text_from_txt(data)


def _extract_text_from_xlsx(data: bytes) -> List[Dict]:
    if pd is None:
        return [_page_record("")]
    try:
        bio = io.BytesIO(data)
        # Read all sheets and concatenate with sheet headers
        xls = pd.ExcelFile(bio)
        pages: List[Dict] = []
        page_num = 1
        for sheet in xls.sheet_names:
            df = xls.parse(sheet)
            df_pages = _extract_pages_from_dataframe(df, table_name=f"sheet_{sheet}")
            # Prefix each page with the sheet name
            for pg in df_pages:
                pg["text"] = f"Sheet: {sheet}\n" + pg["text"]
                pg["page"] = page_num
                pg["page_hash"] = hashlib.md5(pg["text"].encode()).hexdigest()
                pages.append(pg)
                page_num += 1
        return pages if pages else [_page_record("")]
    except Exception:
        return [_page_record("")]


def _infer_sql_type(dtype: str) -> str:
    d = str(dtype).lower()
    if any(k in d for k in ["int64", "int32", "int16", "int8", "uint", "bool"]):
        return "INTEGER"
    if any(k in d for k in ["float", "double"]):
        return "REAL"
    if "datetime" in d or "date" in d:
        return "DATETIME"
    return "TEXT"


def _dataframe_schema_via_sqlite(df: 'pd.DataFrame', table_name: str) -> str:
    # Build schema text from pandas dtypes, and also confirm via sqlite PRAGMA
    try:
        conn = sqlite3.connect(":memory:")
        cols = ", ".join([f"[{c}] {_infer_sql_type(df[c].dtype)}" for c in df.columns])
        conn.execute(f"CREATE TABLE [{table_name}] ({cols});")
        df.to_sql(table_name, conn, if_exists='append', index=False)
        cur = conn.execute(f"PRAGMA table_info([{table_name}]);")
        info = cur.fetchall()
        conn.close()
        schema_cols = ", ".join([f"{row[1]} {row[2]}" for row in info])  # row[1]=name, row[2]=type
        return f"Schema: {table_name}({schema_cols})"
    except Exception:
        # Fallback to pandas dtype mapping only
        schema_cols = ", ".join([f"{c} {_infer_sql_type(df[c].dtype)}" for c in df.columns])
        return f"Schema: {table_name}({schema_cols})"


def _extract_pages_from_dataframe(df: 'pd.DataFrame', table_name: str, max_rows_per_page: int = 200) -> List[Dict]:
    if df is None or df.empty:
        return [_page_record("")]
    schema_line = _dataframe_schema_via_sqlite(df, table_name)
    columns_line = "Columns: " + ", ".join([str(c) for c in df.columns])
    pages: List[Dict] = []
    total_rows = len(df)
    page_number = 1
    for start in range(0, total_rows, max_rows_per_page):
        end = min(start + max_rows_per_page, total_rows)
        chunk = df.iloc[start:end]
        # Convert rows into plain text lines that survive cleaning
        row_lines: List[str] = []
        for idx, row in chunk.iterrows():
            pairs = [f"{col}={str(row[col])}" for col in df.columns]
            row_lines.append("; ".join(pairs))
        text_block = "\n".join([schema_line, columns_line, f"Rows {start+1}-{end} of {total_rows}:"] + row_lines)
        pages.append(_page_record(text_block, page_number))
        page_number += 1
    return pages


def _extract_text_from_image(data: bytes, lang: str = 'ben+eng') -> List[Dict]:
    try:
        img = Image.open(io.BytesIO(data))
        text = pytesseract.image_to_string(enhance_image(img), lang=lang)
        return [_page_record(text, 1)]
    except Exception:
        return [_page_record("")]


def extract_text_from_document(
    file_bytes: bytes,
    file_name: Optional[str] = None,
    lang: str = 'ben+eng',
    dpi: int = 300,
) -> List[Dict]:
    """
    Extract text from various document types, returning a list of page-like records:
    [{"page": int, "text": str, "page_hash": str}].

    Supported by default:
    - .pdf (OCR via pdf2image + Tesseract)
    - .txt/.md (plain text)
    - .docx (python-docx)
    - .html/.htm (BeautifulSoup)
    - .csv (raw text)
    - .png/.jpg/.jpeg/.tiff/.bmp (OCR via Tesseract)
    Unknown extensions fallback to UTF-8 text decoding.
    """
    ext = (os.path.splitext(file_name or "")[1].lower() or "").strip()

    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes, lang=lang, dpi=dpi)
    if ext in {".txt", ".md"}:
        return _extract_text_from_txt(file_bytes)
    if ext == ".docx":
        return _extract_text_from_docx(file_bytes)
    if ext in {".html", ".htm"}:
        return _extract_text_from_html(file_bytes)
    if ext == ".csv":
        return _extract_text_from_csv(file_bytes)
    if ext in {".xlsx"}:  # Excel support via pandas
        return _extract_text_from_xlsx(file_bytes)
    if ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
        return _extract_text_from_image(file_bytes, lang=lang)

    # Fallback: try as plain text
    return _extract_text_from_txt(file_bytes)